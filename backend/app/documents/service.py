"""Moderation pack generation.

The pack is generated *only* from the moderation record: every text block
in the output comes from a recorded field (framework, recommendation,
decision or audit metadata). There is no free-text injection point, so the
pack is provably derived from the record by construction.
"""

import uuid
from dataclasses import dataclass
from datetime import UTC, datetime
from io import BytesIO

from docx import Document
from reportlab.lib.pagesizes import A4
from reportlab.lib.units import mm
from reportlab.pdfgen.canvas import Canvas
from sqlalchemy import func, select
from sqlalchemy.orm import Session

from app.anonymisation.models import AnonymisationMapEntry
from app.audit.recorder import AuditRecorder
from app.core.hashing import content_hash_bytes
from app.framework.models import Criterion, Procurement
from app.ingestion.storage import ObjectStorage
from app.moderation.models import ModerationDecision, ModerationPack
from app.scoring.models import Recommendation, ScoringRun


class PackGenerationError(Exception):
    """The pack cannot be generated; message is safe to show."""


@dataclass(frozen=True)
class PackRow:
    criterion_ref: str
    criterion_title: str
    bidder_token: str
    recommended_score: int | None
    final_score: int
    action: str
    rationale: str | None
    variance: int
    confidence_tier: str
    citation_count: int
    justification: str


@dataclass(frozen=True)
class PackData:
    procurement_title: str
    procurement_reference: str
    lock_hash: str
    model_version: str
    generated_at: str
    rows: list[PackRow]


def build_pack_data(session: Session, procurement: Procurement) -> PackData:
    """Assemble pack content strictly from recorded fields."""
    if procurement.framework_lock_hash is None or procurement.pinned_model_version is None:
        raise PackGenerationError("The framework must be locked before a pack exists.")

    rows_query = session.execute(
        select(Recommendation, ScoringRun, ModerationDecision)
        .join(ScoringRun, ScoringRun.id == Recommendation.run_id)
        .join(
            ModerationDecision,
            ModerationDecision.recommendation_id == Recommendation.id,
        )
        .where(ScoringRun.procurement_id == procurement.id)
    ).all()
    if not rows_query:
        raise PackGenerationError(
            "No moderated decisions exist yet; the pack is generated from the "
            "moderation record."
        )

    criteria = {
        criterion.id: criterion
        for criterion in session.scalars(
            select(Criterion).where(Criterion.procurement_id == procurement.id)
        ).all()
    }
    tokens = {
        entry.bidder_id: entry.token
        for entry in session.scalars(
            select(AnonymisationMapEntry).where(
                AnonymisationMapEntry.procurement_id == procurement.id
            )
        ).all()
    }

    rows: list[PackRow] = []
    for recommendation, run, decision in rows_query:
        criterion = criteria.get(run.criterion_id)
        rows.append(
            PackRow(
                criterion_ref=criterion.ref if criterion else "",
                criterion_title=criterion.title if criterion else "",
                bidder_token=tokens.get(run.bidder_id, "Bidder"),
                recommended_score=recommendation.score,
                final_score=decision.final_score,
                action=decision.action,
                rationale=decision.rationale,
                variance=recommendation.variance,
                confidence_tier=recommendation.confidence_tier,
                citation_count=len(recommendation.citations),
                justification=recommendation.justification,
            )
        )
    rows.sort(key=lambda row: (row.criterion_ref, row.bidder_token))

    return PackData(
        procurement_title=procurement.title,
        procurement_reference=procurement.reference,
        lock_hash=procurement.framework_lock_hash,
        model_version=procurement.pinned_model_version,
        generated_at=datetime.now(UTC).isoformat(timespec="seconds"),
        rows=rows,
    )


def render_docx(data: PackData) -> bytes:
    document = Document()
    document.add_heading("Moderation pack", level=0)
    document.add_paragraph(f"Procurement: {data.procurement_title}")
    document.add_paragraph(f"Reference: {data.procurement_reference}")
    document.add_paragraph(f"Framework lock hash: {data.lock_hash}")
    document.add_paragraph(f"Pinned model version: {data.model_version}")
    document.add_paragraph(f"Generated at: {data.generated_at}")
    document.add_paragraph(
        "Every entry below is derived from the recorded moderation decision "
        "and its underlying recommendation. AI scores, humans moderate, AI "
        "documents."
    )

    for row in data.rows:
        document.add_heading(
            f"{row.criterion_ref} {row.criterion_title} — {row.bidder_token}", level=2
        )
        recommended = (
            str(row.recommended_score)
            if row.recommended_score is not None
            else "escalated (no auto-recommendation)"
        )
        document.add_paragraph(
            f"Recommended score: {recommended} — {row.citation_count} citations — "
            f"variance: {row.confidence_tier} ({row.variance} band(s))"
        )
        document.add_paragraph(
            f"Moderation: {row.action}ed — final score {row.final_score}"
            if row.action == "confirm"
            else f"Moderation: amended — final score {row.final_score}"
        )
        if row.rationale:
            document.add_paragraph(f"Moderator's rationale: {row.rationale}")
        if row.justification:
            document.add_paragraph(f"Recommendation justification: {row.justification}")

    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def render_pdf(data: PackData) -> bytes:
    buffer = BytesIO()
    page = Canvas(buffer, pagesize=A4)
    _, height = A4
    margin = 20 * mm
    y = height - margin

    def line(text: str, size: int = 10) -> None:
        nonlocal y
        if y < margin:
            page.showPage()
            y = height - margin
        page.setFont("Helvetica", size)
        page.drawString(margin, y, text[:110])
        y -= size + 4

    line("Moderation pack", 16)
    line(f"Procurement: {data.procurement_title}", 11)
    line(f"Reference: {data.procurement_reference}")
    line(f"Framework lock hash: {data.lock_hash}")
    line(f"Pinned model version: {data.model_version}")
    line(f"Generated at: {data.generated_at}")
    y -= 6

    for row in data.rows:
        line(f"{row.criterion_ref} {row.criterion_title} — {row.bidder_token}", 12)
        recommended = (
            str(row.recommended_score)
            if row.recommended_score is not None
            else "escalated (no auto-recommendation)"
        )
        line(
            f"Recommended score: {recommended} — {row.citation_count} citations — "
            f"variance: {row.confidence_tier} ({row.variance} band(s))"
        )
        line(f"Moderation: {row.action} — final score {row.final_score}")
        if row.rationale:
            line(f"Moderator's rationale: {row.rationale}")
        y -= 4

    page.save()
    return buffer.getvalue()


def generate_pack(
    session: Session,
    recorder: AuditRecorder,
    storage: ObjectStorage,
    *,
    tenant_schema: str,
    procurement: Procurement,
    file_format: str,
    generated_by: uuid.UUID | None,
) -> ModerationPack:
    if file_format not in ("docx", "pdf"):
        raise PackGenerationError("The pack format must be docx or pdf.")
    data = build_pack_data(session, procurement)
    content = render_docx(data) if file_format == "docx" else render_pdf(data)

    next_version = (
        session.scalar(
            select(func.coalesce(func.max(ModerationPack.version), 0)).where(
                ModerationPack.procurement_id == procurement.id
            )
        )
        or 0
    ) + 1
    pack_id = uuid.uuid4()
    object_key = (
        f"{tenant_schema}/procurements/{procurement.id}/packs/"
        f"moderation-pack-v{next_version}-{pack_id}.{file_format}"
    )
    storage.put(
        object_key,
        content,
        (
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            if file_format == "docx"
            else "application/pdf"
        ),
    )
    pack = ModerationPack(
        id=pack_id,
        procurement_id=procurement.id,
        version=next_version,
        object_key=object_key,
        content_hash=content_hash_bytes(content),
        file_format=file_format,
        generated_by=generated_by,
    )
    session.add(pack)
    session.flush()
    recorder.record(
        "moderation_pack.generated",
        entity_type="moderation_pack",
        entity_id=str(pack.id),
        after_hash=pack.content_hash,
        model_version=procurement.pinned_model_version,
    )
    return pack
